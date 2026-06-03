from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Asus\OneDrive\Documentos\BOOTCAMP\PROYECTOFINAL")
MANUALES = ROOT / "Manuales"
OUT = ROOT / "outputs" / "manual-20260603-playverse" / "manuales"
ASSETS = OUT / "assets"

LOGO = MANUALES / "LOGO.png"
ERD = MANUALES / "RELACIONES Y ENTIDADES DEL PROYECTO.png"

SCREENSHOTS = {
    "home": ASSETS / "playverse-home.png",
    "catalog": ASSETS / "playverse-catalogo.png",
    "login": ASSETS / "playverse-login.png",
    "team": ASSETS / "playverse-equipo.png",
    "contact": ASSETS / "playverse-contacto.png",
    "admin": ASSETS / "playverse-admin.png",
    "admin_catalog": ASSETS / "playverse-admin-catalogo-form.png",
}

TECNICO_OUT = MANUALES / "PlayVerse_Manual_Tecnico.docx"
USUARIO_OUT = MANUALES / "PlayVerse_Manual_Usuario.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(85, 85, 85)
GREEN = RGBColor(21, 145, 82)
PINK = RGBColor(198, 20, 92)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF7EF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_inches: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_run(run, size=None, bold=None, color=None, font="Calibri", italic=None) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def setup_doc(title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header_p = section.header.paragraphs[0]
    header_p.text = title
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header_p.runs[0], size=9, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("PlayVerse | Proyecto final")
    set_run(run, size=9, color=MUTED)
    return doc


def add_cover(doc: Document, manual_name: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        r = p.add_run()
        r.add_picture(str(LOGO), width=Inches(1.55))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PlayVerse")
    set_run(run, size=28, bold=True, color=GREEN)
    set_paragraph_spacing(p, before=14, after=2, line=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(manual_name)
    set_run(run, size=22, bold=True, color=INK)
    set_paragraph_spacing(p, after=6, line=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run(run, size=12, color=MUTED)
    set_paragraph_spacing(p, after=18)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1.65, 4.65])
    rows = [
        ("Sistema", "PlayVerse - Tu universo de videojuegos en un solo lugar"),
        ("Version", "1.0 - Borrador de entrega"),
        ("Fecha", "03 de junio de 2026"),
        ("Equipo", "Juan Buenaño, Jaime Carriel"),
        ("Area", "Bootcamp Full Stack Developer"),
    ]
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for par in cell.paragraphs:
                set_paragraph_spacing(par, after=0)
                for run in par.runs:
                    set_run(run, size=10.5, bold=(cell is row.cells[0]), color=INK)
    doc.add_page_break()


def add_para(doc: Document, text: str, bold=False, color=None, style=None) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run(run, bold=bold, color=color or RGBColor(0, 0, 0))
    set_paragraph_spacing(p)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        set_paragraph_spacing(p, after=4)


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        set_paragraph_spacing(p, after=4)


def add_note(doc: Document, title: str, text: str, fill: str = PALE_GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    run = p.add_run(title + ": ")
    set_run(run, bold=True, color=INK)
    run = p.add_run(text)
    set_run(run, color=INK)
    set_paragraph_spacing(p, after=0)
    doc.add_paragraph()


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], LIGHT_BLUE)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_margins(cell)
            for par in cell.paragraphs:
                set_paragraph_spacing(par, after=0)
                for run in par.runs:
                    set_run(run, size=9.5, bold=(row is table.rows[0]), color=INK)
    doc.add_paragraph()


def add_image(doc: Document, image: Path, caption: str, width=6.2) -> None:
    if not image.exists():
        add_note(doc, "Captura pendiente", f"No se encontro el archivo {image.name}.", fill="FFF2CC")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image), width=Inches(width))
    set_paragraph_spacing(p, after=2, line=1.0)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run(run, size=9, color=MUTED, italic=True)
    set_paragraph_spacing(cap, after=8)


def add_code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run(run, size=9.5, font="Consolas", color=INK)
    set_paragraph_spacing(p, after=0, line=1.15)
    doc.add_paragraph()


def build_tecnico() -> None:
    doc = setup_doc("Manual Tecnico | PlayVerse")
    add_cover(doc, "Manual Tecnico", "Documento para desarrolladores y mantenimiento del sistema")

    doc.add_heading("1. Introduccion y objetivos", level=1)
    add_para(doc, "Este manual describe la arquitectura, instalacion, dependencias y mantenimiento de PlayVerse, una tienda academica de videojuegos construida con Angular, ASP.NET Core y SQL Server.")
    add_bullets(doc, [
        "Guiar a desarrolladores nuevos en la estructura del frontend, backend y base de datos.",
        "Documentar las tecnologias, servicios, controladores y entidades principales.",
        "Definir instrucciones de instalacion, ejecucion local y preparacion para despliegue.",
        "Registrar politicas de autenticacion, roles, permisos y resolucion de errores comunes.",
    ])

    doc.add_heading("2. Glosario tecnico", level=1)
    add_simple_table(doc, ["Termino", "Descripcion"], [
        ["API REST", "Interfaz HTTP que expone recursos como juegos, usuarios, roles y wishlist."],
        ["JWT", "Token firmado usado para identificar al usuario autenticado."],
        ["Guard", "Proteccion de rutas Angular que valida si existe una sesion activa."],
        ["DTO", "Objeto de transferencia de datos entre capas del backend."],
        ["DbContext", "Clase de Entity Framework Core que representa la base SQL Server."],
        ["CORS", "Politica que permite al frontend llamar a la API desde localhost."],
    ], [1.55, 4.95])

    doc.add_heading("3. Requisitos y entorno de desarrollo", level=1)
    add_simple_table(doc, ["Categoria", "Requisito"], [
        ["Sistema operativo", "Windows 10/11 o equivalente con Node.js, .NET SDK y SQL Server."],
        ["Frontend", "Node.js, npm 11.12.1, Angular CLI compatible con Angular 21."],
        ["Backend", ".NET SDK 9, ASP.NET Core, acceso a SQL Server 2022."],
        ["Base de datos", "SQL Server 2022 y script PLAYVERSE_DDL.sql."],
        ["Herramientas", "Visual Studio Code, Postman, Git, navegador moderno."],
    ], [1.6, 4.9])

    doc.add_heading("4. Librerias y frameworks", level=1)
    add_simple_table(doc, ["Capa", "Tecnologias"], [
        ["Frontend", "Angular 21, Angular Material 21, TypeScript 5.9, RxJS, tslib, Vitest."],
        ["Backend", "ASP.NET Core .NET 9, JWT Bearer, OpenAPI, Scalar, Serilog."],
        ["Datos", "Entity Framework Core 9, Microsoft.EntityFrameworkCore.SqlServer, SQL Server 2022."],
        ["Infraestructura", "Dockerfile, render.yaml, variables de entorno, SMTP/Brevo."],
    ], [1.45, 5.05])

    doc.add_heading("5. Arquitectura del sistema", level=1)
    add_para(doc, "PlayVerse sigue una arquitectura separada por responsabilidades. El frontend Angular consume servicios HTTP; la API ASP.NET Core orquesta controladores, servicios de aplicacion, dominio, repositorios y persistencia SQL Server.")
    add_simple_table(doc, ["Capa", "Responsabilidad", "Ubicacion"], [
        ["UI Angular", "Pantallas publicas, login, perfil y panel administrativo.", "FrontEnd/SteamClone/src/app"],
        ["Core Angular", "Constantes, modelos, servicios API, auth guard e interceptor.", "src/app/Core"],
        ["API", "Controladores REST, CORS, middleware de errores y documentacion OpenAPI.", "Steam.Web.Api"],
        ["Aplicacion", "Servicios de Auth, Catalogo, Comunidad, Roles y Usuario.", "SteamApplication"],
        ["Dominio", "Entidades, contexto SQL Server y contratos de repositorio.", "SteamDomain"],
        ["Infraestructura", "Repositorios concretos y UnitOfWork.", "SteamInfrastructure"],
    ], [1.25, 3.0, 2.25])

    doc.add_heading("5.1 Estructura de carpetas del frontend", level=2)
    add_code(doc, """FrontEnd/SteamClone/src/app
|-- Core
|   |-- api              # Servicios HTTP y modelos
|   |-- auth             # AuthService, guard e interceptor
|   |-- constants        # Nombre, slogan, llaves de storage
|-- Features
|   |-- public           # Inicio, juegos, quienes somos, contacto
|   |-- auth             # Login y registro
|   |-- admin            # Panel privado
|-- shared
    |-- layout           # Header, navbar, footer
    |-- ui               # Product card y dialogos""")

    doc.add_heading("5.2 Modulos, componentes y servicios Angular", level=2)
    add_simple_table(doc, ["Elemento", "Descripcion"], [
        ["Home / Products / About / Contact", "Pantallas publicas de navegacion, catalogo, equipo y contacto."],
        ["Login", "Formulario de acceso y registro de usuario."],
        ["AdminDashboard", "Panel privado con secciones de perfil, resumen, catalogo, contenido, tienda, comunidad, accesos y cuenta."],
        ["CatalogService", "Consume endpoints de juegos, generos, developers, publishers, DLC, ofertas y logros."],
        ["StoreCommunityService", "Consume wishlist, libreria, reviews, sesiones, logros y amigos."],
        ["RolesService", "Lista roles, permisos, usuarios y asignaciones."],
        ["AuthService", "Login, registro, renovacion, perfil, cambio y recuperacion de contrasena."],
    ], [1.9, 4.6])

    doc.add_heading("5.3 Backend y endpoints", level=2)
    add_simple_table(doc, ["Controlador", "Responsabilidad"], [
        ["AuthController", "Login, renovacion de token, registro, recuperacion y cambio de contrasena."],
        ["GamesController", "CRUD de juegos y gestion de logros, DLC y ofertas."],
        ["Developers / Publishers / Genres", "Mantenimiento de maestros del catalogo."],
        ["Wishlist / Library", "Wishlist y compras del usuario."],
        ["Reviews / Sessions / Friends / UserAchievements", "Comunidad y actividad del jugador."],
        ["Roles / User", "Roles, permisos, usuarios y cuenta personal."],
    ], [2.1, 4.4])

    doc.add_heading("5.4 Base de datos", level=2)
    add_para(doc, "La base de datos PLAYVERSE contiene entidades para catalogo, usuarios, roles, permisos, actividad de usuario y comunidad. El script principal se encuentra en Base de datos/PLAYVERSE_DDL.sql.")
    add_image(doc, ERD, "Figura 1. Diagrama de relaciones y entidades del proyecto.", width=6.4)

    doc.add_heading("6. Instalacion y ejecucion local", level=1)
    add_steps(doc, [
        "Clonar o copiar el proyecto en el equipo de desarrollo.",
        "Ejecutar el script Base de datos/PLAYVERSE_DDL.sql en SQL Server 2022.",
        "Configurar la cadena de conexion database-1 y variables Jwt/Smtp en appsettings o variables de entorno.",
        "Restaurar y ejecutar el backend desde BackEnd/Steam-Clone/Steam_Backend.",
        "Instalar dependencias del frontend con npm install en FrontEnd/SteamClone.",
        "Levantar Angular con npm run start y abrir http://localhost:4200/.",
    ])
    add_code(doc, """# Backend
cd BackEnd/Steam-Clone/Steam_Backend
dotnet restore
dotnet build
dotnet run --project Steam.Web.Api

# Frontend
cd FrontEnd/SteamClone
npm install
npm run start""")

    doc.add_heading("7. Configuracion", level=1)
    add_simple_table(doc, ["Variable", "Uso"], [
        ["ConnectionStrings__database-1", "Cadena de conexion SQL Server."],
        ["Jwt__Issuer", "Emisor valido del token."],
        ["Jwt__Audience", "Audiencia valida del token."],
        ["Jwt__PrivateKey", "Clave privada para firmar JWT."],
        ["Smtp__Host / Port / User / Password / From", "Configuracion de envio de correos."],
        ["FirstAppTime__User__Email / Password", "Credenciales iniciales del usuario root."],
    ], [2.3, 4.2])

    doc.add_heading("8. Seguridad y accesos", level=1)
    add_bullets(doc, [
        "La autenticacion usa JWT Bearer y el frontend guarda token, refresh token y usuario en localStorage.",
        "El guard de Angular protege la ruta /admin y redirige a /login si no hay sesion activa.",
        "El backend define roles base: admin, developer y user.",
        "Las politicas de autorizacion se construyen con permisos como GAMES_MANAGE, USERS_READ y ROLES_ASSIGN.",
        "No se deben subir credenciales SMTP, JWT ni cadenas de conexion a GitHub.",
    ])

    doc.add_heading("9. Mantenimiento y troubleshooting", level=1)
    add_simple_table(doc, ["Problema", "Causa probable", "Solucion"], [
        ["El catalogo muestra datos de respaldo", "API local apagada o sin conexion.", "Levantar Steam.Web.Api y revisar environment.apiUrl."],
        ["No permite entrar a /admin", "No existe token o usuario activo.", "Iniciar sesion con credenciales validas y revisar localStorage."],
        ["Error de SQL Server", "Cadena de conexion incorrecta o script no ejecutado.", "Validar database-1 y ejecutar PLAYVERSE_DDL.sql."],
        ["Error SMTP", "Variables SMTP incompletas.", "Configurar Smtp__Host, User, Password y From."],
        ["CORS bloquea llamadas", "Origen no permitido.", "Agregar origen en FrontendCorsPolicy."],
    ], [1.7, 2.25, 2.55])

    doc.add_heading("10. Anexos", level=1)
    add_bullets(doc, [
        "Coleccion Postman: BackEnd/Steam-Clone/Steam_Backend/PlayVerse_Full_Demo.postman_collection.json.",
        "Ambiente Postman: API.postman_environment.template.json.",
        "Despliegue: render.yaml y Dockerfile.",
        "Logo y diagramas: carpeta Manuales.",
    ])
    doc.save(TECNICO_OUT)


def build_usuario() -> None:
    doc = setup_doc("Manual de Usuario | PlayVerse")
    add_cover(doc, "Manual de Usuario", "Guia para usuarios finales y operadores administrativos")

    doc.add_heading("1. Introduccion", level=1)
    add_para(doc, "PlayVerse es una tienda academica de videojuegos que permite explorar juegos, consultar informacion publica, iniciar sesion y operar un panel privado para administrar catalogo, contenido y usuarios segun el rol asignado.")
    add_note(doc, "Derechos de autor", "Documento elaborado para fines academicos del proyecto final PlayVerse. El uso de marcas o imagenes de videojuegos dentro de las capturas corresponde al material de demostracion del sistema.")

    doc.add_heading("2. Glosario", level=1)
    add_simple_table(doc, ["Termino", "Explicacion sencilla"], [
        ["Catalogo", "Listado de juegos disponibles en la tienda."],
        ["Wishlist", "Lista personal de juegos que el usuario desea guardar."],
        ["Libreria", "Juegos adquiridos o asociados al usuario."],
        ["Rol", "Nivel de permiso del usuario, por ejemplo admin, developer o user."],
        ["Panel administrativo", "Zona privada para gestionar informacion del sistema."],
        ["Sesion", "Estado en el que el usuario ya inicio sesion."],
    ], [1.5, 5.0])

    doc.add_heading("3. Guia de inicio rapido", level=1)
    add_steps(doc, [
        "Abrir el navegador y entrar a http://localhost:4200/ cuando el sistema este corriendo.",
        "Revisar la pantalla de inicio para identificar las opciones principales.",
        "Entrar a Juegos para consultar el catalogo y filtrar por categoria.",
        "Usar Login para ingresar o crear una cuenta.",
        "Si el usuario tiene rol autorizado, entrar a Admin para operar el panel privado.",
    ])
    add_image(doc, SCREENSHOTS["home"], "Figura 1. Pantalla de inicio de PlayVerse.", width=6.35)

    doc.add_heading("4. Acceso al sistema", level=1)
    add_para(doc, "La pantalla de login permite iniciar sesion con un usuario existente y tambien crear una cuenta nueva si el sistema tiene habilitado el registro.")
    add_image(doc, SCREENSHOTS["login"], "Figura 2. Pantalla de login y acceso al sistema.", width=6.35)
    doc.add_heading("4.1 Iniciar sesion", level=2)
    add_steps(doc, [
        "Presionar el boton Login en la parte superior derecha.",
        "Ingresar correo y contrasena en el formulario.",
        "Presionar Ingresar al panel o el boton equivalente.",
        "Verificar que aparezca el saludo del usuario y las opciones Perfil/Admin/Salir.",
    ])
    doc.add_heading("4.2 Cerrar sesion", level=2)
    add_steps(doc, [
        "Presionar Salir en la barra superior.",
        "Confirmar que el sistema vuelva a mostrar la opcion Login.",
    ])

    doc.add_heading("5. Secciones principales", level=1)
    doc.add_heading("5.1 Inicio", level=2)
    add_para(doc, "Muestra la identidad del sistema, el slogan y accesos rapidos hacia el catalogo o inicio de sesion.")
    doc.add_heading("5.2 Juegos", level=2)
    add_para(doc, "Permite consultar el catalogo de videojuegos. Si la API no esta disponible, el frontend puede mostrar un catalogo de respaldo para mantener la demostracion.")
    add_image(doc, SCREENSHOTS["catalog"], "Figura 3. Catalogo publico con filtros de categoria.", width=6.35)
    doc.add_heading("5.3 Quienes somos", level=2)
    add_para(doc, "Presenta informacion del proyecto y del equipo responsable.")
    add_image(doc, SCREENSHOTS["team"], "Figura 4. Seccion de equipo e informacion del proyecto.", width=6.35)
    doc.add_heading("5.4 Contactanos", level=2)
    add_para(doc, "Muestra datos de contacto y formulario para consultas o soporte.")
    add_image(doc, SCREENSHOTS["contact"], "Figura 5. Seccion de contacto.", width=6.35)

    doc.add_heading("6. Panel administrativo", level=1)
    add_para(doc, "El panel administrativo esta disponible para usuarios autenticados con permisos. Desde alli se gestionan catalogo, contenido, tienda, comunidad, accesos y cuenta.")
    add_image(doc, SCREENSHOTS["admin"], "Figura 6. Vista superior del panel administrativo.", width=6.35)
    add_note(doc, "Nota de captura", "La captura muestra el panel con sesion local de demostracion. Si el backend esta apagado, aparece el mensaje 'No se pudo completar la operacion'.")

    doc.add_heading("6.1 Listar registros", level=2)
    add_steps(doc, [
        "Entrar al panel Admin.",
        "Seleccionar la seccion Catalogo o el modulo correspondiente.",
        "Presionar Actualizar catalogo para solicitar datos actuales al backend.",
        "Revisar la lista o tabla de registros disponibles.",
    ])
    doc.add_heading("6.2 Crear un juego", level=2)
    add_image(doc, SCREENSHOTS["admin_catalog"], "Figura 7. Formulario de creacion de juego y maestros del catalogo.", width=6.35)
    add_steps(doc, [
        "En el panel, seleccionar Catalogo.",
        "Completar los campos Titulo, Descripcion, Imagen del juego, Fecha de lanzamiento y Precio.",
        "Seleccionar Desarrollador, Publisher y Generos si estan disponibles.",
        "Presionar Crear juego.",
        "Verificar el mensaje de exito y actualizar el catalogo.",
    ])
    doc.add_heading("6.3 Editar un registro", level=2)
    add_steps(doc, [
        "Ubicar el juego o maestro que se desea modificar.",
        "Presionar la accion Editar del registro.",
        "Cambiar los campos necesarios.",
        "Presionar Guardar cambios.",
        "Confirmar que la informacion actualizada aparezca en el listado.",
    ])
    doc.add_heading("6.4 Eliminar un registro", level=2)
    add_steps(doc, [
        "Ubicar el registro dentro del listado.",
        "Presionar Eliminar.",
        "Leer el mensaje de confirmacion.",
        "Confirmar la eliminacion solo si el registro ya no debe existir.",
        "Actualizar el listado para validar que desaparecio.",
    ])
    add_note(doc, "Precaucion", "Eliminar registros puede afectar relaciones como wishlist, libreria, reviews, DLC u ofertas. En una demostracion, usar datos de prueba.")

    doc.add_heading("7. Gestion de perfil y cuenta", level=1)
    add_bullets(doc, [
        "Perfil permite revisar datos del usuario autenticado.",
        "Cuenta permite cambiar contrasena o usar flujos de seguridad si estan habilitados.",
        "El usuario debe cerrar sesion al terminar una demostracion o uso compartido.",
    ])

    doc.add_heading("8. Preguntas frecuentes", level=1)
    add_simple_table(doc, ["Pregunta", "Respuesta"], [
        ["No veo datos en el catalogo administrativo.", "Revise que el backend y SQL Server esten encendidos. Si la API esta apagada, el panel mostrara error."],
        ["Puedo entrar al catalogo publico sin login?", "Si. Las pantallas Inicio, Juegos, Quienes somos y Contactanos son publicas."],
        ["Por que no aparece Admin?", "El usuario no inicio sesion o no tiene rol/permisos suficientes."],
        ["Puedo crear registros sin datos maestros?", "Algunos campos pueden depender de developers, publishers y generos cargados en la API."],
        ["Que hago si olvide mi contrasena?", "Usar el flujo de recuperacion si esta configurado el servicio SMTP."],
    ], [2.15, 4.35])

    doc.add_heading("9. Contacto y soporte", level=1)
    add_bullets(doc, [
        "Correo: contacto@playverse.dev",
        "Telefono: +593 99 610 1033",
        "Equipo: Juan Buenaño y Jaime Carriel",
    ])
    doc.save(USUARIO_OUT)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build_tecnico()
    build_usuario()
    print(TECNICO_OUT)
    print(USUARIO_OUT)
